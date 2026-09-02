#!/usr/bin/env python3
"""Build the eight EDI progress reports.

The supplied example .docx is used as the template, so the school header, the
logo, the instructor line, the fonts and the page setup all carry over exactly.
Only the body is replaced.

usage: make_reports.py <template.docx> <output_dir> <shots_dir>
"""
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import report_content as C

TEMPLATE, OUT, SHOTS = (Path(a) for a in sys.argv[1:4])
OUT.mkdir(parents=True, exist_ok=True)

BODY_W = 6.9                    # A4 minus the template's 0.5in margins
MAX_H = 6.5                     # so a tall screenshot cannot eat a whole page
GREY = RGBColor(0x59, 0x59, 0x59)
STATUS_INK = {"Completed": RGBColor(0x1E, 0x7A, 0x3C),
              "In Progress": RGBColor(0xB0, 0x6A, 0x00),
              "Not Started": RGBColor(0x59, 0x59, 0x59),
              "Blocked": RGBColor(0xB3, 0x26, 0x1E)}


def clear_body(doc):
    """Empty the template's body but keep its section properties."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def para(doc, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align:
        p.alignment = align
    return p


def run(p, text, bold=False, size=11, colour=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if colour:
        r.font.color.rgb = colour
    return r


def labelled(doc, label, text):
    p = para(doc)
    run(p, label + " ", bold=True)
    run(p, text)


def heading(doc, text):
    p = para(doc, space_after=4)
    p.paragraph_format.space_before = Pt(12)
    run(p, text, bold=True, size=13)


def cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tcPr.append(borders)


def shade(cell, hexcolour):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolour)
    tcPr.append(shd)


def milestone_table(doc, index):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(4.1), Inches(1.5), Inches(1.3))

    head = table.rows[0].cells
    for cell, text, w in zip(head, ("Task / Milestone", "Assigned Member", "Status"), widths):
        cell.width = w
        shade(cell, "F2F2F2")
        cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run(p, text, bold=True, size=10)

    for task, owner, codes in C.TASKS:
        status = C.STATUS[codes[index]]
        cells = table.add_row().cells
        for cell, text, w in zip(cells, (task, owner, status), widths):
            cell.width = w
            cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if text is status:
                run(p, text, size=10, bold=True, colour=STATUS_INK[status])
            else:
                run(p, text, size=10)
    return table


def figure(doc, path, caption, number):
    w, h = Image.open(path).size
    width = min(BODY_W, MAX_H * w / h)
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
    cap = para(doc, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(cap, f"Figure {number}. ", bold=True, size=9, colour=GREY)
    run(cap, caption, size=9, colour=GREY, italic=True)


def build(index):
    doc = docx.Document(str(TEMPLATE))
    clear_body(doc)
    spec = C.REPORTS[index]

    title = para(doc, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(title, C.TITLE, bold=True, size=14)

    labelled(doc, "Date:", C.DATES[index])
    labelled(doc, "Team / Group Name:", C.GROUP)
    labelled(doc, "Team Members:", C.TEAM)

    heading(doc, "1. Project Overview")
    labelled(doc, "Innovation Goal:", C.GOAL)
    labelled(doc, "Target Audience / End User:", C.AUDIENCE)

    heading(doc, "2. Key Milestones & Completed Tasks")
    milestone_table(doc, index)
    para(doc, space_after=0)

    heading(doc, "3. Current Focus & Key Wins")
    labelled(doc, "What we worked on recently:", spec["recent"])
    labelled(doc, "Key achievements or insights:", spec["wins"])

    heading(doc, "4. Challenges & Obstacles")
    labelled(doc, "Current Bottlenecks:", spec["bottleneck"])
    labelled(doc, "Proposed Solution / Action Plan:", spec["solution"])
    labelled(doc, "Help or Resources Needed:", spec["help"])

    heading(doc, "5. Next Steps")
    # the template carries no list styles, so the numbers are written in
    for n, step in enumerate(spec["next"], start=1):
        p = para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run(p, f"{n}.  ", bold=True)
        run(p, step)

    heading(doc, "6. Evidence")
    p = para(doc, space_after=10)
    run(p, "Screenshots taken from the running system on this date.", italic=True, colour=GREY, size=10)
    for n, (filename, caption) in enumerate(spec["shots"], start=1):
        figure(doc, SHOTS / C.SHOT_DIRS[index] / filename, caption, n)

    day = C.DATES[index].split("/")[0]
    name = f"EDI Progress Report - AI Cashier System (G12IS) - {int(day):02d} Aug 2026.docx"
    path = OUT / name
    doc.save(str(path))
    return path


for i in range(len(C.DATES)):
    p = build(i)
    print(f"  {p.name}   {p.stat().st_size // 1024} KB")
